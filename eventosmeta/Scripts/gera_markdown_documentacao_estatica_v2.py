import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import re
from docx import Document
from datetime import datetime
from pathlib import Path
import os

class DocGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Documentação ESTÁTICO")
        self.root.geometry("800x600")

        self.word_path = tk.StringVar()
        self.code_path = tk.StringVar()

        # Frame para seleções
        frame_select = ttk.Frame(root, padding="10")
        frame_select.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        ttk.Label(frame_select, text="Arquivo Word (ESTATICO_*.docx):").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(frame_select, textvariable=self.word_path, width=60).grid(row=0, column=1, padx=5)
        ttk.Button(frame_select, text="Selecionar", command=self.select_word).grid(row=0, column=2)

        ttk.Label(frame_select, text="Arquivo Código Django (.txt):").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(frame_select, textvariable=self.code_path, width=60).grid(row=1, column=1, padx=5)
        ttk.Button(frame_select, text="Selecionar", command=self.select_code).grid(row=1, column=2)

        # Botão gerar
        ttk.Button(frame_select, text="Gerar Markdown", command=self.generate_md).grid(row=2, column=0, columnspan=3, pady=20)

        # Log
        ttk.Label(root, text="Logs:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        self.log_text = scrolledtext.ScrolledText(root, height=20, width=90)
        self.log_text.grid(row=2, column=0, padx=10, pady=5, sticky=(tk.W, tk.E, tk.N, tk.S))

        root.columnconfigure(0, weight=1)
        root.rowconfigure(2, weight=1)

    def log(self, msg):
        self.log_text.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')} - {msg}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def select_word(self):
        file = filedialog.askopenfilename(
            title="Selecione arquivo Word",
            filetypes=[("DOCX", "*.docx")]
        )
        if file and 'ESTATICO_' not in Path(file).name:
            messagebox.showwarning("Aviso", "O arquivo deve iniciar com ESTATICO_")
            return
        self.word_path.set(file)
        self.log(f"Word selecionado: {file}")

    def select_code(self):
        file = filedialog.askopenfilename(
            title="Selecione arquivo código",
            filetypes=[("TXT", "*.txt")]
        )
        self.code_path.set(file)
        self.log(f"Código selecionado: {file}")

    def extract_from_word(self, doc_path):
        try:
            doc = Document(doc_path)
            sections = {'adrs': [], 'regras': [], 'glossario': [], 'mer': []}
            current_section = None

            keywords = {
                'adrs': ['ADR', 'DECISÃO ARQUITETURAL', 'ARCHITECTURE DECISION'],
                'regras': ['REGRAS DE NEGÓCIO', 'REGRAS IMUTÁVEIS'],
                'glossario': ['GLOSSÁRIO'],
                'mer': ['MER', 'MODELO ENTIDADE RELACIONAMENTO', 'ENTIDADE-RELACIONAMENTO']
            }

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                text_upper = text.upper()
                for sec, kws in keywords.items():
                    if any(kw in text_upper for kw in kws):
                        current_section = sec
                        break
                if current_section and text:
                    sections[current_section].append(text)

            mer_tables = []
            for table in doc.tables:
                md_table = self.table_to_md(table)
                if md_table:
                    mer_tables.append(md_table)

            self.log("Extração do Word concluída.")
            return sections['adrs'], sections['regras'], sections['glossario'], mer_tables
        except Exception as e:
            raise ValueError(f"Erro ao extrair Word: {str(e)}")

    def table_to_md(self, table):
        if not table.rows:
            return ""
        md_lines = []
        header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        md_lines.append('| ' + ' | '.join(header) + ' |')
        md_lines.append('| ' + ' --- |' * len(header) + ' |')
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_lines.append('| ' + ' | '.join(cells) + ' |')
        return '\n'.join(md_lines)

    def extract_django_code(self, code_path):
        try:
            with open(code_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Apps from imports
            apps = list(set(re.findall(r'from\s+(\w+)\.models\s+import', content)))

            # Models
            models_info = []
            model_matches = re.finditer(r'class\s+(\w+)\s*\([^)]*models\.Model[^)]*\)\s*:\s*(.*?)(?=class\s+\w+|(?:def\s+|$))', content, re.DOTALL | re.IGNORECASE)
            for match in model_matches:
                name = match.group(1)
                body = match.group(2)
                fields = re.findall(r'(\w+)\s*=\s*models\.(\w+)\s*\(', body)
                indexes = re.findall(r'indexes\s*=\s*\[([^\]]*)\]', body, re.DOTALL)
                models_info.append({'name': name, 'fields': fields, 'indexes': indexes[0] if indexes else ''})

            # Services
            services = re.findall(r'class\s+(\w+Service)\s*\(', content)

            # URLs
            urls = re.findall(r"path\s*\(\s*['\"]([^'\"]+)['\"]\s*,\s*([^),]+)", content)

            # Management commands
            commands = re.findall(r'class\s+(\w+Command)\s*\(BaseCommand\):', content)

            # Admin actions
            admin_actions = re.findall(r'def\s+(\w+)_action\s*\(self,', content)

            arch = {
                'apps': apps,
                'models': models_info,
                'services': services,
                'urls': urls,
                'commands': commands,
                'admin_actions': admin_actions
            }
            self.log("Extração do código concluída.")
            return arch
        except Exception as e:
            raise ValueError(f"Erro ao extrair código: {str(e)}")

    def build_md(self, adrs, regras, glossario, mer_tables, arch):
        date_str = datetime.now().strftime('%Y%m%d')
        md = f"# ESTÁTICO_Documentacao-Consolidada_{date_str}\n\n"

        # 1. Arquitetura Geral
        md += "## 1. Arquitetura Geral do Sistema\n\n"
        if arch['apps']:
            md += f"### Apps Django: {', '.join(arch['apps'])}\n\n"
        if arch['services']:
            md += "### Services:\n"
            for s in arch['services']:
                md += f"- {s} [VALIDAR]\n"
            md += "\n"
        if arch['urls']:
            md += "### URLs:\n"
            for path, view in arch['urls']:
                md += f"- `{path}` -> {view.strip()} [REVISAR]\n"
            md += "\n"
        if arch['commands']:
            md += "### Management Commands:\n"
            for c in arch['commands']:
                md += f"- {c} [VALIDAR]\n"
            md += "\n"
        if arch['admin_actions']:
            md += "### Admin Actions:\n"
            for a in arch['admin_actions']:
                md += f"- {a} [CONFLITO]\n"
            md += "\n"
        md += "### Models:\n\n"
        for model in arch['models']:
            md += f"#### {model['name']}\n"
            for field, ftype in model['fields']:
                md += f"- `{field}`: {ftype} [VALIDAR]\n"
            if model['indexes']:
                md += f"**Indexes:** {model['indexes']} [REVISAR]\n\n"
            else:
                md += "\n"

        # 2. ADRs
        md += "## 2. Decisões Arquiteturais (ADRs)\n\n"
        for adr in adrs:
            md += f"- {adr} [REVISAR]\n"
        md += "\n"

        # 3. Regras
        md += "## 3. Regras de Negócio Imutáveis\n\n"
        for regra in regras:
            md += f"- {regra} [VALIDAR]\n"
        md += "\n"

        # 4. Glossário
        md += "## 4. Glossário Técnico e Funcional\n\n"
        for g in glossario:
            md += f"- {g} [REVISAR]\n"
        md += "\n"

        # 5. MER
        md += "## 5. Modelo Entidade-Relacionamento (MER)\n\n"
        md += "### Tabelas MER do Word:\n\n"
        for table_md in mer_tables:
            md += f"{table_md}\n\n"
        md += "### Referência aos Models (ver seção 1)\n\n"

        return md

    def generate_md(self):
        try:
            word = self.word_path.get()
            code = self.code_path.get()
            if not word or not code:
                raise ValueError("Selecione ambos os arquivos.")

            self.log("Iniciando extração...")
            adrs, regras, glossario, mer_tables = self.extract_from_word(word)
            arch = self.extract_django_code(code)

            self.log("Gerando Markdown...")
            md_content = self.build_md(adrs, regras, glossario, mer_tables, arch)

            date_str = datetime.now().strftime('%Y%m%d')
            out_dir = Path('docs/estática')
            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / f'ESTÁTICO_Documentacao-Consolidada_{date_str}.md'

            with open(out_path, 'w', encoding='utf-8') as f:
                f.write(md_content)

            self.log(f"Arquivo salvo: {out_path}")
            messagebox.showinfo("Sucesso", f"Documentação gerada em:\n{out_path}")
        except Exception as e:
            self.log(f"ERRO: {str(e)}")
            messagebox.showerror("Erro", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = DocGenerator(root)
    root.mainloop()

    