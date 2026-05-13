import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from pathlib import Path
import re
import datetime

# Diretório base fixo
base_dir = Path(r"C:\PMS\PMS2025\Inscr-Meta")

# Função auxiliar para construir texto de parágrafo com formatação Markdown
def build_para_text(para):
    parts = []
    for run in para.runs:
        text = run.text
        if text.strip():
            if run.bold and run.italic:
                parts.append(f"***{text}***")
            elif run.bold:
                parts.append(f"**{text}**")
            elif run.italic:
                parts.append(f"*{text}*")
            else:
                parts.append(text)
    return ''.join(parts)

# Função para converter tabela DOCX para Markdown
def build_table_md(table):
    if not table.rows:
        return ''
    num_cols = len(table.rows[0].cells)
    # Cabeçalho
    header_row = [cell.text.strip() for cell in table.rows[0].cells]
    rows_md = ['| ' + ' | '.join(header_row) + ' |']
    # Separador
    separator = '|' + ' --- |' * num_cols
    rows_md.append(separator)
    # Linhas restantes
    for row in table.rows[1:]:
        row_cells = [cell.text.strip() for cell in row.cells]
        rows_md.append('| ' + ' | '.join(row_cells) + ' |')
    return '\n'.join(rows_md)

# Função principal para converter DOCX para Markdown
def docx_to_markdown(doc):
    md = []
    # Parágrafos
    for para in doc.paragraphs:
        text = build_para_text(para)
        if not text.strip():
            continue
        style = para.style.name.lower()
        if 'heading' in style:
            level_match = re.search(r'(\d+)', style)
            level = int(level_match.group(1)) if level_match else 1
            md.append('#' * level + ' ' + text)
        elif any(word in style for word in ['list bullet', 'list number', 'list', 'bullet']):
            md.append('- ' + text)
        else:
            md.append(text)
    # Tabelas
    for table in doc.tables:
        table_md = build_table_md(table)
        if table_md:
            md.append(table_md)
    return '\n\n'.join(md)

# Função para converter arquivo DOCX para Markdown
def convert_docx_to_md(docx_path, md_path):
    doc = Document(docx_path)
    md_content = docx_to_markdown(doc)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

# Função para extrair seção específica do Markdown
def get_section(content, title):
    pattern = rf'(?m)^#{{2}}\s+{re.escape(title)}\s*\n((?:^[^#]{{2}}.*| ^\s*$|\n)*)'
    match = re.search(pattern, content, re.DOTALL | re.MULTILINE)
    if match:
        return match.group(1).rstrip()
    return ''

# Função para extrair models do código
def extract_models(code_text):
    pattern = r'class\s+(?P<name>\w*(?:Model|model)\w*)\s*\((?P<base>[^)]+)\)\s*:\s*(?P<body>[\s\S]*?)(?=\nclass\s+\w|\Z)' 
    models = []
    for match in re.finditer(pattern, code_text, re.DOTALL | re.MULTILINE):
        name = match.group('name')
        base = match.group('base').strip()
        body = match.group('body').strip()
        models.append(f"### {name} (herda de {base})\n\n")

